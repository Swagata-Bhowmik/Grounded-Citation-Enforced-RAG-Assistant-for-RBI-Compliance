"""
build_project_bible.py
======================
Generates 'Grounded_Project_Bible.docx' — a complete, beginner-friendly,
interview-prep guide to the whole project: concepts taught from zero, the full
build story, the code, deployment/git, and a large Q&A bank.

Rich formatting: colored headings, shaded callout "cards", code blocks, Q&A
cards, tables, emojis. Run:  python build/build_project_bible.py
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "docs" / "Grounded_Project_Bible.docx"

# ---- palette -------------------------------------------------------------
NAVY   = RGBColor(0x12, 0x14, 0x3A)
BLUE   = RGBColor(0x15, 0x65, 0xC0)
PURPLE = RGBColor(0x6A, 0x1B, 0x9A)
GREEN  = RGBColor(0x2E, 0x7D, 0x32)
TEAL   = RGBColor(0x00, 0x83, 0x8F)
ORANGE = RGBColor(0xEF, 0x6C, 0x00)
PINK   = RGBColor(0xC2, 0x18, 0x5B)
RED    = RGBColor(0xC6, 0x28, 0x28)
GREY   = RGBColor(0x44, 0x44, 0x44)
BLACK  = RGBColor(0x1A, 0x1A, 0x2E)

doc = Document()
# base style
_st = doc.styles["Normal"]
_st.font.name = "Calibri"
_st.font.size = Pt(11)
_st.font.color.rgb = BLACK


# ---- low-level helpers ---------------------------------------------------
def _shade(el, hex_color):
    """Apply background shading to a cell or paragraph element."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    el.append(shd)


def _add_runs(p, text, base_color=None, base_bold=False, base_size=None):
    """Add text to paragraph, honoring **bold** and `code` inline markers."""
    import re
    # split on **bold** and `code`
    tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        bold = base_bold
        mono = False
        val = tok
        if tok.startswith("**") and tok.endswith("**"):
            bold = True
            val = tok[2:-2]
        elif tok.startswith("`") and tok.endswith("`"):
            mono = True
            val = tok[1:-1]
        r = p.add_run(val)
        r.bold = bold
        if base_color is not None:
            r.font.color.rgb = base_color
        if base_size is not None:
            r.font.size = Pt(base_size)
        if mono:
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x66)
    return p


def H1(text, color=PURPLE):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = color
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]) if isinstance(color, tuple) else "6A1B9A")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def H2(text, color=BLUE):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = color
    return p


def H3(text, color=TEAL):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = color
    return p


def P(text, size=11, color=None, italic=False, align=None):
    p = doc.add_paragraph()
    _add_runs(p, text, base_color=color, base_size=size)
    if italic:
        for r in p.runs:
            r.italic = True
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def BULLETS(items, color=None):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        _add_runs(p, it, base_color=color)
    return


def NUMBERED(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        _add_runs(p, it)
    return


def CALLOUT(title, body, fill="EDE7F6", bar=PURPLE, emoji=""):
    """A shaded single-cell 'card' with a colored title line."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    _shade(cell._tc.get_or_add_tcPr(), fill)
    # left color bar via table borders
    tp = cell.paragraphs[0]
    r = tp.add_run(f"{emoji} {title}".strip())
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = bar
    if body:
        bp = cell.add_paragraph()
        _add_runs(bp, body)
    # tighten
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def CODE(text):
    """A shaded monospace code block."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    _shade(cell._tc.get_or_add_tcPr(), "1E1E2E")
    for i, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xE6, 0xE6, 0xF0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


_QN = [0]
def QA(q, a, tips=None):
    _QN[0] += 1
    n = _QN[0]
    # question line
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    _shade(cell._tc.get_or_add_tcPr(), "E8F0FE")
    p = cell.paragraphs[0]
    r = p.add_run(f"Q{n}. ")
    r.bold = True; r.font.color.rgb = BLUE; r.font.size = Pt(11)
    _add_runs(p, q, base_color=NAVY)
    for rr in p.runs:
        rr.bold = True
    # answer
    ap = doc.add_paragraph()
    ar = ap.add_run("A. ")
    ar.bold = True; ar.font.color.rgb = GREEN
    _add_runs(ap, a)
    if tips:
        tp = doc.add_paragraph()
        _add_runs(tp, "**💡 Say this too:** " + tips)
        for rr in tp.runs:
            rr.font.size = Pt(10); rr.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def TABLE(headers, rows, header_fill="6A1B9A"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i]._tc.get_or_add_tcPr(), header_fill)
        p = hdr[i].paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            _add_runs(cells[i].paragraphs[0], str(val))
            for rr in cells[i].paragraphs[0].runs:
                rr.font.size = Pt(10)
    doc.add_paragraph()
    return t


def PAGEBREAK():
    doc.add_page_break()


def SPACER(pts=6):
    doc.add_paragraph().paragraph_format.space_after = Pt(pts)


# ======================================================================
#  COVER
# ======================================================================
def cover():
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("🤖 GROUNDED"); r.bold = True; r.font.size = Pt(40); r.font.color.rgb = PURPLE
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Citation-Enforced Agentic RAG Assistant\nfor RBI Banking Compliance")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = BLUE
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("📘 THE COMPLETE PROJECT GUIDE & INTERVIEW BIBLE")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = GREEN
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Everything explained from zero — concepts, code, deployment, and a full interview Q&A bank")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Author: Swagata Bhowmik  ·  MSc Data Science")
    r.font.size = Pt(12); r.font.color.rgb = NAVY; r.bold = True
    PAGEBREAK()


cover()
print("framework ready")


# ======================================================================
#  PART 1 — START HERE
# ======================================================================
H1("📑 What's inside this document")
P("This guide is built so you can walk in cold and walk out able to defend every "
  "line of this project. Read it top to bottom once, then use the Q&A bank to revise.")
TABLE(["Part", "What it covers"], [
    ["1. Start Here", "How to use this, and the elevator pitches (30-sec, 2-min)"],
    ["2. Vocabulary from Zero", "Every technical word explained simply — AI, LLM, embeddings, RAG, NLI, CI…"],
    ["3. The Build Story", "Each stage of the project: what, why, how (with code), what we found"],
    ["4. The Codebase", "Every file and what it does, in plain terms"],
    ["5. Deployment & GitHub", "Exact commands, secrets safety, cloud deploy, the real problems we hit"],
    ["6. Results & Limitations", "The numbers, the key decisions, and the honest edges"],
    ["7. Interview Q&A Bank", "100+ questions with answers, grouped by topic"],
    ["8. Cheat Sheet", "The numbers and one-liners to memorize the night before"],
], header_fill="1565C0")

H2("🧭 How to use this")
BULLETS([
    "**First read:** Parts 1–3 slowly. Don't skip Part 2 even if a word looks scary — every term is explained.",
    "**Before an interview:** re-read Part 8 (cheat sheet) and skim Part 7 (Q&A).",
    "**If you forget a concept:** search Part 2 for the word.",
    "**Golden rule for interviews:** it is 100%% fine to say *\"let me think\"* and explain in your own words. "
    "This doc gives you the words; the understanding comes from reading it a couple of times.",
])

CALLOUT("The single most important sentence to remember",
        "Grounded answers banking-compliance questions from real RBI regulations, cites the exact page it "
        "used, and refuses to answer when the documents don't support one — instead of making something up.",
        fill="E8F5E9", bar=GREEN, emoji="⭐")

H2("🎤 The 30-second elevator pitch")
P("\"I built **Grounded**, an AI assistant that answers banking-compliance questions over real RBI "
  "regulations. The twist is trust: every answer **cites the exact source page**, and if the regulations "
  "don't contain the answer, it **declines instead of hallucinating**. Under the hood it uses "
  "**retrieval-augmented generation** with hybrid search, a re-ranker, a citation-enforcement rule, and a "
  "small decision-making agent — plus an evaluation set and an automated quality gate. It runs for free.\"")

H2("🎤 The 2-minute version")
P("\"Compliance teams at banks must answer precise questions against dense RBI rulebooks, and a confidently "
  "wrong answer is a real regulatory risk. A normal chatbot will happily hallucinate. Grounded takes the "
  "opposite approach.")
P("It reads 10 real RBI Master Directions, cuts them into small page-tagged pieces, and turns each piece "
  "into a numeric 'meaning fingerprint'. When you ask a question, it searches two ways at once — by keyword "
  "and by meaning — merges the results, then a second, more careful model re-ranks the shortlist. If the best "
  "piece of evidence isn't strong enough, the system refuses. If it is, a language model writes an answer "
  "grounded only in that evidence, citing the page.")
P("Then I proved it works: I wrote a graded question set, measured retrieval and refusal accuracy, and added "
  "a faithfulness check that flags any invented claim. A GitHub Actions pipeline re-runs those checks on every "
  "code change and blocks anything that lowers quality. Finally I deployed it as a free public web app.\"")
PAGEBREAK()

# ======================================================================
#  PART 2 — VOCABULARY FROM ZERO
# ======================================================================
H1("🧠 Part 2 — Every Technical Word, Explained Simply", color=BLUE)
P("Read this once and the rest of the document (and most AI interviews) will make sense. Each term has a "
  "plain-English meaning, a simple analogy, and where it shows up in **our** project.")

def TERM(word, plain, analogy, ours):
    H3("🔹 " + word)
    P("**What it means:** " + plain)
    P("**Analogy:** " + analogy)
    P("**In Grounded:** " + ours)
    SPACER(2)

H2("Group A — The AI basics")
TERM("AI (Artificial Intelligence)",
     "Software that does tasks we'd normally call 'intelligent' — understanding language, recognising images, making decisions.",
     "A very well-read intern who can read fast and summarise, but needs clear instructions and checking.",
     "Our whole assistant is an AI system; but note we deliberately keep it on a short leash (cite or refuse).")
TERM("Machine Learning (ML)",
     "A way of building software that **learns patterns from examples** instead of being programmed with fixed rules.",
     "Instead of writing rules for 'what a cat looks like', you show the computer 10,000 cat photos and it learns.",
     "The models we use (for embeddings, re-ranking, faithfulness, generation) were all trained by ML on huge text datasets.")
TERM("Model",
     "The trained 'brain' file that an ML system produces — a big set of numbers (weights) that turns an input into an output.",
     "A recipe learned from tasting thousands of dishes; give it ingredients (input), it predicts a dish (output).",
     "We use several ready-made models: bge-small (embeddings), a cross-encoder (re-ranking), an NLI model (fact-checking), and an LLM (writing answers).")
TERM("Training vs Inference",
     "**Training** = the slow, one-time process of learning from data. **Inference** = actually using the trained model to get an answer (fast).",
     "Training is studying for years; inference is answering one exam question.",
     "We do **no training** — we only do inference with pre-trained models. That's why it runs on a normal laptop.")

H2("Group B — Language models")
TERM("LLM (Large Language Model)",
     "A very large ML model trained to predict the next word, which makes it able to write and answer in fluent language.",
     "Autocomplete on superpowers — it has read much of the internet and continues text sensibly.",
     "An LLM writes the final answer, but ONLY from the evidence we hand it. We use a small local one (Qwen2.5-3B) and Google Gemini in the cloud.")
TERM("Token",
     "A chunk of text the model actually reads — usually a word or part of a word (about ¾ of a word on average).",
     "Lego bricks of text: 'compliance' might be one brick, 'KYC' another.",
     "We size our document pieces in tokens (384 each) because the embedding model has a hard limit of 512 tokens.")
TERM("Parameters",
     "The individual numbers inside a model that were tuned during training. More parameters ≈ more capacity (and more RAM).",
     "The knobs on a giant mixing board; '3B' means 3 billion knobs.",
     "Our local LLM 'Qwen2.5-3B' has 3 billion parameters — small enough for a laptop, big enough to write a clean cited answer.")
TERM("Quantization",
     "Shrinking a model by storing its numbers with less precision (e.g. 4 bits instead of 16), so it uses far less memory with little quality loss.",
     "Saving a photo as a slightly smaller JPEG — much smaller file, looks almost identical.",
     "Our local model is a 4-bit 'quantized GGUF' file, which is why a 3-billion-parameter model fits and runs on CPU.")
TERM("Prompt / System prompt",
     "The instruction you give an LLM. The **system prompt** is the standing rulebook; the user prompt is the actual question.",
     "The system prompt is the job description; the user prompt is today's specific task.",
     "Our system prompt (in prompts.yaml) hard-codes the rules: answer only from context, cite the page, or refuse.")
TERM("Temperature",
     "A dial (0–1+) controlling randomness. Low = focused and repeatable; high = creative and varied.",
     "Low temperature is a careful accountant; high is a brainstorming poet.",
     "We set it to 0.1 — a compliance assistant must be factual and near-deterministic, never 'creative'.")
TERM("Hallucination",
     "When an LLM states something false but confident, because it's predicting plausible text, not checking facts.",
     "A student who didn't study but writes a confident, wrong exam answer.",
     "Eliminating hallucination is the entire point of Grounded — we ground answers in retrieved text and refuse when unsure.")

H2("Group C — Search & meaning")
TERM("NLP (Natural Language Processing)",
     "The field of AI focused on understanding and generating human language.",
     "Teaching computers to read and write, not just calculate.",
     "The whole project is applied NLP: reading regulations, understanding questions, writing answers.")
TERM("Embedding (vector)",
     "A list of numbers that captures the **meaning** of a piece of text. Similar meanings get similar number-lists.",
     "A GPS coordinate for meaning — sentences about the same topic land near each other on the map.",
     "We turn every document piece into a 384-number embedding so we can find relevant pieces by meaning, not just keywords.")
TERM("Dimension",
     "How many numbers are in each embedding (here, 384). Each dimension captures some aspect of meaning.",
     "A description using 384 sliders instead of just 'height and weight'.",
     "Our bge-small model outputs 384-dimensional vectors — a good balance of quality and speed.")
TERM("Cosine similarity",
     "A score (−1 to 1) for how close two embeddings point in the same direction — i.e. how similar in meaning.",
     "Two arrows: pointing the same way = 1 (identical meaning); at right angles = 0 (unrelated).",
     "We rank document pieces by cosine similarity to the question. In our test: 0.978 to a restatement, 0.55 to unrelated text.")
TERM("Semantic search",
     "Searching by meaning (using embeddings) rather than exact words — so 'car' can match 'automobile'.",
     "A librarian who understands what you mean, not just the exact words you said.",
     "One half of our retrieval. Great for paraphrased questions, weaker on exact codes/numbers.")
TERM("Keyword search / BM25",
     "Searching by exact words. **BM25** is the classic scoring formula: it rewards rare matching words and isn't fooled by long documents.",
     "Ctrl+F with a smart scorekeeper that weighs rare, important words more.",
     "The other half of our retrieval. Essential here because compliance users search exact clause numbers and acronyms like 'KFS'.")
TERM("TF-IDF (background for BM25)",
     "An older formula: a word matters more if it appears **often in this document** (TF) but is **rare across all documents** (IDF). BM25 is an improved version.",
     "'KYC' appearing a lot in one doc and rarely elsewhere makes it a strong signal for that doc.",
     "BM25 (which builds on TF-IDF ideas) is our keyword scorer.")

H2("Group D — The RAG machinery")
TERM("RAG (Retrieval-Augmented Generation)",
     "The core pattern: first **retrieve** relevant text from your documents, then let the LLM **generate** an answer using that text. So answers are based on real sources, not the model's memory.",
     "Open-book exam: look up the right page first, then answer using it — instead of guessing from memory.",
     "Grounded IS a RAG system. Retrieval finds the regulation text; generation writes the cited answer from it.")
TERM("Chunking",
     "Cutting big documents into small pieces so we can retrieve just the relevant bit, not a whole 300-page PDF.",
     "Turning a book into labelled index cards, one idea per card.",
     "We chunk each PDF into 384-token pieces with 64-token overlap, and tag every chunk with its page number.")
TERM("Overlap",
     "Letting consecutive chunks share some text, so an idea sitting on a boundary isn't split and lost.",
     "Overlapping roof tiles so nothing falls through the gaps.",
     "64 tokens of overlap between our chunks.")
TERM("Tokenizer",
     "The tool that splits text into tokens. Each model has its own.",
     "The specific way one person slices bread vs another.",
     "We measure chunk size with the embedding model's OWN tokenizer, so no chunk is secretly too big for it.")
TERM("Vector database / ChromaDB",
     "A special database that stores embeddings and, given a query embedding, instantly finds the closest ones (nearest-neighbour search).",
     "A filing cabinet organised by meaning that hands you the closest cards in milliseconds.",
     "We use ChromaDB to store each chunk's vector + text + (document, page) metadata. The metadata is what enables citations.")
TERM("Nearest-neighbour search",
     "Finding the items whose vectors are closest to a query vector.",
     "Finding the nearest coffee shops to your location on a map.",
     "How semantic retrieval pulls the most meaning-similar chunks for a question.")
TERM("Reciprocal Rank Fusion (RRF)",
     "A simple, robust way to merge two ranked lists (e.g. keyword and semantic results) into one, using only the **rank** of each item, not its raw score.",
     "Two judges each rank the contestants; RRF combines their rankings fairly without needing their scoring scales to match.",
     "We fuse BM25 and semantic results with RRF, avoiding the problem that their scores live on totally different scales.")
TERM("Re-ranking / Cross-encoder",
     "A second, more accurate model that reads the **question and a candidate passage together** and scores true relevance. Slower, so used only on a shortlist.",
     "A careful second reader who re-checks the top 50 shortlisted cards with the question in hand.",
     "A ms-marco cross-encoder re-ranks our ~50 fused candidates. Its score also becomes our confidence signal for refusing.")
TERM("Bi-encoder vs Cross-encoder",
     "A **bi-encoder** embeds query and passage separately then compares (fast, used for first-pass search). A **cross-encoder** reads them together (accurate, used for re-ranking).",
     "Bi-encoder: two people describe a movie separately and you compare notes. Cross-encoder: one person watches both together and judges the match.",
     "Our embedder (bge-small) is a bi-encoder; our re-ranker is a cross-encoder. Classic two-stage retrieval.")

H2("Group E — The agent, trust & quality")
TERM("Agent",
     "An AI setup that **decides what to do** (which step, which tool, whether to answer) rather than always following one fixed path.",
     "A junior analyst who decides: do I need to look this up? is my evidence enough? should I answer or say I don't know?",
     "Our agent decides answer-vs-decline, then verifies the answer actually contains a citation.")
TERM("LangGraph / State machine",
     "A tool for building agents as a **graph of steps** with decision branches. A state machine = a system that moves between defined states along allowed transitions.",
     "A flowchart the program actually follows: box → decision diamond → box.",
     "Our flow: retrieve → assess → (answer | decline) → verify → cite. The answer/decline branch is the agent's real decision.")
TERM("Citation / Grounding",
     "**Grounding** = every claim is backed by supplied evidence. **Citation** = pointing to exactly which source (here, which page) backs it.",
     "A well-referenced essay where every claim has a footnote.",
     "Our answers cite the exact PDF and page, e.g. [02_Digital_Lending_Directions_2025.pdf p.8].")
TERM("Citation enforcement / Refusal",
     "A rule that blocks answering unless the evidence is strong enough; otherwise the system declines.",
     "A cautious expert who says 'I don't have that' rather than guessing.",
     "If the top re-rank score is below a threshold (0.0), we return a fixed refusal message. This is our safety feature.")
TERM("NLI (Natural Language Inference)",
     "A model that decides whether a claim is **entailed by** (follows from), **neutral to**, or **contradicts** a piece of text.",
     "A fact-checker: 'does this evidence actually prove this sentence?'",
     "We score faithfulness by checking each answer sentence against the evidence with an NLI model — offline, no cost.")
TERM("Faithfulness",
     "Whether the answer is truly supported by the retrieved evidence (didn't add anything extra).",
     "Marking an open-book answer: did they use the book, or sneak in made-up facts?",
     "Grounded answer scored 0.97 faithfulness; a deliberately fake answer scored 0.01 — a clean separation.")
TERM("Golden set / Ground truth",
     "A hand-made set of questions with known-correct answers, used to measure how good the system is.",
     "An answer key for a quiz.",
     "We wrote 26 questions (24 answerable, mapped to their correct document; 2 out-of-scope to test refusal).")
TERM("Recall@k",
     "Of the times the right document exists, how often it appears in the top **k** retrieved results. Recall@5 = in the top 5.",
     "Did the right book show up in your top-5 search results?",
     "Our retrieval recall@5 = 1.00 across 24 questions (the right document was always in the top 5).")
TERM("Precision vs Recall",
     "**Precision** = of what you returned, how much was right. **Recall** = of all the right things, how much you found.",
     "Precision = few false alarms; recall = missed nothing.",
     "We report recall@5 for retrieval and accuracy for the refusal decision.")
TERM("Accuracy",
     "The fraction of decisions that were correct.",
     "Marks out of total.",
     "Refusal accuracy = 1.00: every in-scope question was answered and every out-of-scope one declined, correctly.")

H2("Group F — Shipping it")
TERM("Git / GitHub / Repository",
     "**Git** tracks changes to your code (version control). **GitHub** hosts it online. A **repository (repo)** is one project's folder of tracked code.",
     "Git = 'save points' with history; GitHub = the cloud drive that shares them.",
     "The whole project lives in a GitHub repo; each phase was a commit.")
TERM("Commit / Push / Branch",
     "**Commit** = save a snapshot with a message. **Push** = upload commits to GitHub. **Branch** = a parallel line of work; 'main' is the primary one.",
     "Commit = save the draft; push = send it to the shared drive; branch = a separate copy to experiment on.",
     "We committed phase-by-phase with clear messages and pushed to the 'main' branch.")
TERM("CI/CD (Continuous Integration)",
     "Automation that runs checks (like tests) automatically on every code change, catching problems early.",
     "A robot inspector that re-checks the product every time you tweak it.",
     "A GitHub Actions workflow re-runs our quality metrics on every push and fails the build if they drop.")
TERM("Quality gate",
     "A pass/fail bar in CI: if quality drops below a threshold, the build fails and the change is blocked.",
     "A bouncer who won't let a change in if it lowers the score.",
     "Ours requires recall ≥ 0.85 and refusal accuracy ≥ 0.85, or the build turns red.")
TERM("API / API key",
     "An **API** lets one program use another over the internet. An **API key** is your private password to use it.",
     "API = a waiter taking your order to the kitchen; the key = your membership card.",
     "The cloud version uses Google Gemini's API for generation; the API key is kept secret (never in the code).")
TERM("Streamlit",
     "A Python tool that turns a script into an interactive web app with almost no web coding.",
     "PowerPoint for data apps — write Python, get buttons and charts.",
     "Our dashboard and live demo are Streamlit apps.")
TERM("Deployment",
     "Putting your app on a server so anyone with the link can use it.",
     "Moving from cooking at home to opening the restaurant to the public.",
     "We deployed the app free on Streamlit Community Cloud.")
TERM("llama.cpp / GGUF",
     "**llama.cpp** runs LLMs efficiently on CPU. **GGUF** is the compact model file format it uses.",
     "A fuel-efficient engine (llama.cpp) and the special fuel it takes (GGUF).",
     "Our offline LLM is a GGUF Qwen model run by llama.cpp — no GPU, no API.")
PAGEBREAK()
print("part 1+2 done")


# ======================================================================
#  PART 3 — THE BUILD STORY
# ======================================================================
H1("🏗️ Part 3 — How the Project Was Built, Stage by Stage", color=GREEN)
P("The project follows a **3-phase blueprint**: Phase 1 = a working RAG pipeline; Phase 2 = production "
  "quality (hybrid search, re-ranking, citation enforcement); Phase 3 = agentic + shippable (agent, "
  "evaluation, faithfulness, CI, deployment). Below, each stage is explained the same way: **what it is, "
  "why we did it, how we did it (with real code), and what we found.**")

def STAGE(num, title, what, why, how_bullets, code=None, code_expl=None, found=None, interview=None):
    H2(f"{num}. {title}")
    P("**🎯 What it is:** " + what)
    P("**❓ Why we did it:** " + why)
    P("**🔧 How we did it:**")
    BULLETS(how_bullets)
    if code:
        P("**Key code:**", color=NAVY)
        CODE(code)
    if code_expl:
        P("**What that code does:** " + code_expl)
    if found:
        CALLOUT("What we found", found, fill="FFF8E1", bar=ORANGE, emoji="🔎")
    if interview:
        CALLOUT("Interview angle", interview, fill="E3F2FD", bar=BLUE, emoji="🎤")
    SPACER(4)

STAGE(1, "Choosing & profiling the data (the corpus)",
    "The set of real documents the assistant is allowed to read — 10 RBI Master Directions on lending, "
    "credit and customer protection.",
    "A RAG system is only as good as its documents. We needed real, current, machine-readable PDFs on a "
    "coherent theme — and we had to CHECK them before trusting them (the 'data is sacred' rule).",
    ["Downloaded 10 real RBI Master Directions directly from rbi.org.in",
     "**Profiled** every PDF before building anything: page count, characters, and whether text is real "
     "(born-digital) or scanned images needing OCR",
     "Recorded the results in a manifest CSV so the check is reproducible, not a one-off glance"],
    code="import fitz  # PyMuPDF\n"
         "doc = fitz.open(pdf_path)\n"
         "for i in range(doc.page_count):\n"
         "    text = doc[i].get_text('text')\n"
         "    # count chars; flag pages with almost no text (scanned or blank)",
    code_expl="We open each PDF with PyMuPDF and pull the raw text per page, counting characters to spot "
              "scanned or empty pages.",
    found="All 10 documents are born-digital (0 scanned files) — no OCR needed. A 'multi-column' flag turned "
          "out to be a false alarm, and 3 near-empty pages in the NBFC doc were just 'Withdrawn' dividers. "
          "Total: 10 docs, 720 pages, ~1.4M characters.",
    interview="If asked 'how did you ensure data quality?', say: I profiled every PDF for scanned-vs-real text "
              "and character counts BEFORE building the pipeline, and investigated every anomaly rather than "
              "assuming the data was clean.")

STAGE(2, "Chunking the documents",
    "Cutting each PDF into small, page-tagged text pieces (chunks).",
    "You can't hand an LLM a 300-page PDF. Retrieval needs small, focused pieces, and each piece must "
    "remember its page so we can cite it.",
    ["Split text into **384-token** chunks with **64-token overlap**",
     "Measured size in the **embedding model's own tokenizer**, so no chunk exceeds its 512-token limit",
     "Tagged every chunk with its source file and **exact page number** (the foundation of citations)",
     "**Experimented** with 256 / 384 / 480 token sizes and chose 384 with evidence, not by guessing"],
    code="splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(\n"
         "    tokenizer, chunk_size=384, chunk_overlap=64,\n"
         "    separators=['\\n\\n', '\\n', '. ', ' ', ''])\n"
         "# split page-by-page so each chunk keeps its page number",
    code_expl="A recursive splitter tries to break on paragraph, then line, then sentence boundaries — so "
              "chunks stay coherent instead of cut mid-word. We run it per page to preserve page tags.",
    found="384/64 produced **1,252 chunks** (avg ~270 tokens) — the sweet spot between too-fragmented (256) "
          "and too-coarse (480). Zero chunks exceeded the 512 limit.",
    interview="Why 384 tokens? Small enough to give precise citations and fit the embedder, large enough to "
              "hold a complete idea. I validated it with a size experiment, not intuition.")

STAGE(3, "Embedding the chunks",
    "Turning every chunk into a 384-number vector that captures its meaning.",
    "So we can find relevant chunks by **meaning**, not just exact words — which handles paraphrased questions.",
    ["Model: **BAAI/bge-small-en-v1.5** — small, CPU-friendly, fully offline, 384-dim",
     "**L2-normalized** the vectors so cosine similarity is a clean dot-product",
     "Prepended bge's recommended **query instruction** to questions (that's how the model was trained)",
     "Embedded all 1,252 chunks on CPU — no GPU, no API cost"],
    code="model = SentenceTransformer('BAAI/bge-small-en-v1.5', device='cpu')\n"
         "vectors = model.encode(texts, normalize_embeddings=True)  # (1252, 384)",
    code_expl="SentenceTransformer loads the model and turns a list of texts into a matrix of normalized "
              "vectors, ready for similarity comparison.",
    found="Sanity check worked perfectly: a KYC question scored 0.978 similarity to a near-identical "
          "restatement, 0.721 to a same-topic sentence, and only 0.546 to unrelated priority-sector text.",
    interview="Why a small model? For extractive retrieval you don't need a giant embedder; bge-small is "
              "fast, free, offline, and its quality is validated on our own data.")

STAGE(4, "Storing vectors (ChromaDB)",
    "Putting the embeddings into a vector database for instant meaning-based lookup.",
    "We need to find the closest chunks to a question in milliseconds, and keep the page metadata attached.",
    ["Used **ChromaDB** with cosine distance, persisted to disk (survives restarts)",
     "Stored each chunk's **vector + text + metadata (document, page)** together",
     "Passed in **our own** embeddings so the SAME model embeds chunks and queries (avoids a classic RAG mismatch bug)"],
    code="collection.add(ids=ids, embeddings=vectors,\n"
         "               documents=texts, metadatas=metas)\n"
         "hits = collection.query(query_embeddings=[qvec], n_results=5)",
    code_expl="We add all chunks with their metadata, then query with a question's vector to get the nearest chunks.",
    found="Meaning search returned the right KYC pages (38, 6, 40) in milliseconds, each carrying its page for citation.",
    interview="Note: in the final app we actually keep the vectors in a NumPy array + BM25 index for the hybrid "
              "retriever; ChromaDB is the Phase-1 store. Either way the concept — nearest-neighbour by meaning — is the same.")

STAGE(5, "Hybrid retrieval (keyword + meaning)",
    "Searching two ways at once — BM25 keyword search AND semantic search — then merging the results.",
    "Semantic search misses exact tokens (clause numbers, 'KFS'); keyword search misses paraphrases. "
    "Regulation is full of exact numbers, so we need both. **This is why this corpus was chosen** — it "
    "genuinely justifies hybrid retrieval.",
    ["Built a **BM25** keyword index over all 1,252 chunks",
     "Tokenized so clause numbers like `13.03` stay intact as single tokens",
     "Ran BM25 and semantic search in parallel and fused with **Reciprocal Rank Fusion (RRF)**"],
    code="sem = semantic_rank(query, top=50)   # by meaning\n"
         "kw  = bm25_rank(query, top=50)       # by keyword\n"
         "fused = rrf_fuse([sem, kw])          # merge by rank",
    code_expl="Each retriever returns a ranked list of chunk indices; RRF merges them into one order using "
              "ranks (so BM25's and cosine's different score scales don't clash).",
    found="Exact-acronym queries ('KFS') and clause numbers are caught by BM25; paraphrases by semantic search. "
          "The fusion covers both without tuning weights.",
    interview="Explain RRF simply: it merges ranked lists by position, not raw score, so it's robust when the "
              "two searches score on totally different scales.")

STAGE(6, "Cross-encoder re-ranking",
    "A second, more careful model re-scores the shortlist by reading the question and each passage together.",
    "First-pass search is fast but rough. A cross-encoder is far more accurate at judging true relevance — "
    "too slow for all 1,252 chunks, but perfect for the ~50 shortlisted candidates.",
    ["Model: **cross-encoder/ms-marco-MiniLM-L-6-v2**",
     "Two-stage design: hybrid search narrows 1,252 → ~50, then the cross-encoder re-ranks those to a top-k",
     "Reused the cross-encoder's **score as a confidence signal** for the refusal decision"],
    code="pairs  = [(query, chunk_text) for chunk_text in candidates]\n"
         "scores = cross_encoder.predict(pairs)   # higher = more relevant\n"
         "top    = sort_desc(scores)[:top_k]",
    code_expl="The cross-encoder takes (question, passage) pairs and outputs one relevance score each; we keep the highest.",
    found="For the KFS question the true answer (Digital Lending **p.8**) scored **+4.95**, while the next "
          "candidates fell below +0.3 — a decisive top result.",
    interview="Bi-encoder vs cross-encoder is a favourite question: bi-encoder embeds separately (fast, first pass); "
              "cross-encoder reads together (accurate, re-rank). Using both is 'two-stage retrieval'.")

STAGE(7, "Citation enforcement (the refusal rule)",
    "The rule that makes the system decline when evidence is weak, instead of guessing.",
    "In compliance, a confident wrong answer is worse than 'not found'. Refusal is the single most important "
    "trustworthy behaviour.",
    ["If the top cross-encoder score is below a **threshold (0.0)**, return a fixed **refusal message**",
     "Threshold and all prompts live in a **versioned config file** (prompts.yaml), not hard-coded",
     "0.0 is the cross-encoder's natural boundary (its logit 0 ≈ 50%% relevance probability)"],
    code="if top_score < config['answerability_threshold']:  # 0.0\n"
         "    return refusal_message\n"
         "else:\n"
         "    build_cited_context_and_answer()",
    code_expl="A single gate: below-threshold evidence → decline; otherwise assemble the cited context and answer.",
    found="Answerable questions scored well above 0 (up to +7.5); out-of-scope questions topped out at −8.7 and "
          "−9.6 — a huge, safe margin between answer and decline.",
    interview="Why threshold in a config file? So a risk/compliance team can tune strictness without touching code, "
              "and every change is tracked in Git. We even bumped the prompt v1→v2 to fix citation formatting with zero code change.")

STAGE(8, "The LangGraph agent",
    "A small decision-making flow that ties retrieval, the refusal rule, generation, and a citation check together.",
    "Rather than one fixed straight line, an agent DECIDES: is the evidence enough? answer or decline? and then "
    "verifies it actually cited a page. This makes behaviour auditable.",
    ["Graph: **retrieve → assess → (answer | decline) → verify → cite**",
     "The conditional answer/decline edge is the genuine 'agentic' decision",
     "A **verify** node confirms the written answer really contains a page citation",
     "Provider-swappable LLM: the same graph works with the local model or a hosted one"],
    code="g.add_edge(START, 'retrieve')\n"
         "g.add_edge('retrieve', 'assess')\n"
         "g.add_conditional_edges('assess', route, {'generate':'generate','decline':'decline'})\n"
         "g.add_edge('generate', 'verify')",
    code_expl="LangGraph wires named steps into a graph; `route` reads the assessment and sends flow to either "
              "generate or decline.",
    found="The cake question stops at 'assess' and declines — the LLM is never even called (saving time). The "
          "KFS question flows all the way to a verified, cited answer.",
    interview="Why call it an agent and not a function? Because of the conditional decision (answer vs decline) and "
              "the self-check (verify). The graph also makes it easy to add tools later without rewriting the flow.")

STAGE(9, "Generating the answer (the LLM)",
    "The language model that writes the final answer — using ONLY the retrieved, cited evidence.",
    "We need fluent, accurate answers that stay strictly grounded in the evidence and cite pages.",
    ["Local default: **Qwen2.5-3B-Instruct**, 4-bit quantized GGUF, run by **llama.cpp** on CPU — offline, $0",
     "Cloud: **Google Gemini** free tier (so it fits a free host with little RAM)",
     "Both sit behind ONE interface `chat(system, user)`, so swapping providers changes only one file",
     "Temperature 0.1 for factual, near-deterministic answers"],
    code="def get_llm(provider=None):\n"
         "    # auto: use Gemini if an API key is set, else the local model\n"
         "    if provider == 'gemini': return GeminiLLM()\n"
         "    return LocalLLM()",
    code_expl="A factory picks the provider. The agent doesn't know or care which LLM it got — that's the "
              "'provider-swappable' design the brief required.",
    found="A small 3B model is entirely enough for extractive, cited answers — we didn't need a giant frontier model.",
    interview="Why a small/local model? For grounded extraction the LLM mostly rephrases supplied text and cites — "
              "that doesn't need a huge model. Small = free, offline, private (crucial for a bank).")

STAGE(10, "Evaluation (the golden set)",
    "A graded quiz with an answer key, to measure the system objectively.",
    "'It ran without error' is not 'it works'. We measure retrieval and refusal against known-correct answers.",
    ["Hand-authored **26 questions**: 24 answerable (each mapped to its correct document) + 2 out-of-scope",
     "Metric 1 — **retrieval recall@5**: is the right document in the top 5?",
     "Metric 2 — **refusal accuracy**: answer in-scope, decline out-of-scope",
     "Both are **deterministic and LLM-free**, so they're perfect for an automated gate"],
    code="res = evaluate_all(retriever, cfg, golden, k=5)\n"
         "# -> recall_at_k = 1.00 (24 Qs), refusal_accuracy = 1.00 (26 Qs)",
    code_expl="We run every golden question through retrieval and the refusal gate and compare to the known answers.",
    found="**Recall@5 = 1.00** and **refusal accuracy = 1.00**. Honest caveat: the set is small and recall is "
          "document-level — a strong baseline, not an inflated claim.",
    interview="Always volunteer the caveat (small, document-level). Interviewers respect honesty about limitations "
              "far more than a suspiciously perfect story.")

STAGE(11, "Faithfulness scoring (NLI)",
    "An automatic fact-checker that verifies the answer didn't invent anything beyond the evidence.",
    "Even with the right evidence, an LLM can slip in an unsupported detail. We catch that.",
    ["Model: **nli-distilroberta** — decides entailment / neutral / contradiction",
     "Each answer sentence is checked against the context sentences for **entailment (support)**",
     "Fully **offline and free** — no expensive 'LLM-as-judge'"],
    code="score = nli_entailment(answer_sentence, evidence_sentences)\n"
         "# grounded answer -> 0.97 ; fabricated answer -> 0.01",
    code_expl="For each sentence of the answer we ask the NLI model whether the evidence entails it; we average the support.",
    found="We caught a real bug: feeding whole passages made scores mushy; switching to **sentence-level** premises "
          "fixed it. Grounded answer = 0.97, fabricated = 0.01 — a sharp separation.",
    interview="This shows evaluation depth: I didn't just trust the metric, I debugged it (sentence-level vs "
              "passage-level) until it cleanly separated true from invented answers.")

STAGE(12, "CI/CD quality gate",
    "Automation that re-runs the metrics on every code change and blocks regressions.",
    "So quality can't silently break as the project evolves — the difference between a demo and a maintained system.",
    ["`scripts/quality_gate.py` computes the metrics and **exits non-zero** if they fall below thresholds",
     "`tests/test_quality_gate.py` asserts recall ≥ 0.85 and refusal ≥ 0.85",
     "`.github/workflows/quality-gate.yml` runs it on **every push/PR** on a clean Linux machine"],
    code="RECALL_AT_5_MIN = 0.85\nREFUSAL_ACCURACY_MIN = 0.85\n"
         "passed = recall >= RECALL_AT_5_MIN and refusal >= REFUSAL_ACCURACY_MIN\n"
         "sys.exit(0 if passed else 1)   # non-zero fails the CI build",
    code_expl="If either metric drops below the bar, the script exits with an error code, which turns the GitHub build red.",
    found="It rebuilds the retriever from the committed PDFs on a fresh machine — proving true reproducibility, "
          "not 'works on my laptop'.",
    interview="Why only the LLM-free metrics in CI? They're fast and deterministic, so they gate reliably. "
              "Faithfulness needs generated answers (slow on CPU) so it's run locally/in the notebook.")

STAGE(13, "Dashboard, notebook & deployment",
    "The ways a human actually sees and uses the project.",
    "A portfolio piece must be explorable by non-coders and reachable by a public link.",
    ["A guided **Jupyter notebook** tells the full build story cell by cell",
     "A **Streamlit app** gives an interactive 'ask a question' demo with visible citations",
     "A standalone **HTML dashboard** (this walkthrough) needs no server, hosted on **GitHub Pages**",
     "The live app is deployed **free** on **Streamlit Community Cloud** using the Gemini free tier"],
    found="We originally targeted Hugging Face Spaces, but in 2026 HF made compute Spaces paid. We pivoted: "
          "swapped the local LLM for Gemini's free API so the app fits Streamlit's free tier — a real-world "
          "adaptation to a platform change.",
    interview="This pivot is a great story: when the free hosting changed, I used the provider-swappable design I'd "
              "already built to move the LLM to a free API and deploy elsewhere — no rewrite needed.")
PAGEBREAK()
print("part 3 done")


# ======================================================================
#  PART 4 — THE CODEBASE, FILE BY FILE
# ======================================================================
H1("📂 Part 4 — The Codebase, File by File", color=TEAL)
P("If an interviewer opens your repo and asks 'what does this file do?', here's your answer for each one.")

TABLE(["File", "What it does (plain English)"], [
    ["regulatory_corpus/*.pdf", "The 10 real RBI Master Directions — the source data."],
    ["regulatory_corpus/corpus_manifest.csv", "The profiling results: pages, characters, text-quality flags per PDF."],
    ["src/corpus_profiler.py", "Reads each PDF and records quality stats — the 'validate the data first' step."],
    ["src/document_chunker.py", "Cuts PDFs into 384-token, page-tagged chunks using the embedder's tokenizer."],
    ["src/embedder.py", "Loads bge-small and turns text into 384-dim normalized vectors (chunks & queries)."],
    ["src/vector_store.py", "The ChromaDB wrapper: store vectors + metadata, query nearest neighbours."],
    ["src/hybrid_retriever.py", "BM25 + semantic search, RRF fusion, and cross-encoder re-ranking."],
    ["src/citation_guard.py", "The refusal rule + builds the cited context; loads prompts.yaml."],
    ["src/agent.py", "The LangGraph agent: retrieve → assess → answer/decline → verify → cite."],
    ["src/llm_provider.py", "Swappable LLM: LocalLLM (Qwen/llama.cpp) and GeminiLLM behind one chat() interface."],
    ["src/evaluation.py", "Runs the golden set: computes recall@5 and refusal accuracy."],
    ["src/faithfulness.py", "The NLI faithfulness scorer (sentence-level entailment)."],
    ["prompts/prompts.yaml", "Versioned system prompt, answer template, refusal message, and the threshold."],
    ["evaluation/golden_questions.json", "The 26 hand-authored test questions + expected behaviour."],
    ["evaluation/eval_results.json", "The saved metric results the dashboard reads."],
    ["scripts/quality_gate.py", "The CI script: compute metrics, exit non-zero if below threshold."],
    ["tests/test_quality_gate.py", "pytest assertions the CI runs."],
    [".github/workflows/quality-gate.yml", "GitHub Actions config that runs the gate on every push."],
    ["dashboard/streamlit_app.py", "The interactive web app (local LLM or Gemini) with the live demo."],
    ["dashboard/build_html_dashboard.py", "Generates the standalone HTML walkthrough dashboard."],
    ["config/project_paths.py", "Keeps big model/cache files off the full C: drive; portable across machines."],
    ["corpus_embeddings.npy", "Precomputed embeddings committed so the cloud app starts instantly."],
    ["requirements.txt", "Lean, cloud-deployable dependencies."],
    ["requirements-lock.txt", "The full pinned local environment (for exact reproduction)."],
], header_fill="00838F")

CALLOUT("The mental model to remember",
        "Data (PDFs) → profile → chunk → embed → store → **retrieve (hybrid + rerank)** → **guard (cite or "
        "refuse)** → **agent** → **LLM writes cited answer**. Around that core sit evaluation, faithfulness, "
        "the CI gate, and the app/deploy layer.",
        fill="EDE7F6", bar=PURPLE, emoji="🧩")
PAGEBREAK()

# ======================================================================
#  PART 5 — DEPLOYMENT & GITHUB PLAYBOOK
# ======================================================================
H1("🚀 Part 5 — Deployment & GitHub Playbook", color=ORANGE)
P("Everything you did to ship it, and the exact commands — so you can explain (and repeat) the process.")

H2("Git & GitHub — the commands used")
P("**Git** saves versioned snapshots of your code; **GitHub** hosts them online. The typical loop is "
  "**stage → commit → push**.")
CODE("git init                       # start tracking a folder (once)\n"
     "git add .                      # stage all changes for the next snapshot\n"
     "git status                     # see what will be committed\n"
     'git commit -m "message"        # save a snapshot with a description\n'
     "git branch -M main             # name the primary branch 'main'\n"
     "git remote add origin <url>    # link to the GitHub repo (once)\n"
     "git push -u origin main        # upload commits to GitHub\n"
     "\n"
     "# day-to-day after that:\n"
     "git add <files>                # stage specific files\n"
     'git commit -m "what changed"\n'
     "git push                       # send to GitHub")
P("**We committed phase by phase** with meaningful messages (e.g. 'Add free Gemini-backed cloud deploy'). "
  "That history itself tells the project story.")

H2("Keeping secrets & big files out of Git")
P("A `.gitignore` file lists things Git should never upload. Ours excludes the API key file, caches, and "
  "large model binaries.")
CODE(".env\n.streamlit/secrets.toml     # the real API key — NEVER committed\n"
     ".grounded_cache/            # downloaded models\n*.bin\n*.pt")
CALLOUT("Golden rule about the API key",
        "The Gemini API key is a password. It lives only in a git-ignored local file "
        "(.streamlit/secrets.toml) and, in the cloud, in Streamlit's Secrets box. It is never typed into "
        "code or committed. We even keep a `.example` template so others know the format without seeing the key.",
        fill="FFEBEE", bar=RED, emoji="🔒")

H2("Two requirements files, and why")
BULLETS([
    "**requirements.txt** — a lean, CPU-only list (no llama.cpp, no Jupyter) that installs cleanly on the free "
    "cloud host. Streamlit Community Cloud reads this.",
    "**requirements-lock.txt** — the full pinned local environment, for exact reproduction on a dev machine.",
])

H2("Deploying the live app (Streamlit Community Cloud)")
NUMBERED([
    "Push the repo to GitHub.",
    "Go to share.streamlit.io and sign in with GitHub.",
    "Create app → pick the repo, branch `main`, main file `dashboard/streamlit_app.py`.",
    "In Advanced settings → Secrets, paste `GEMINI_API_KEY = \"...\"` (never in code).",
    "Deploy. First build installs deps and downloads the small models; then it's live at a public URL.",
])

H2("Getting the free Gemini API key")
NUMBERED([
    "Go to aistudio.google.com/app/apikey and sign in with Google.",
    "Click Create API key. Copy the long string starting 'AIza…'.",
    "Put it in .streamlit/secrets.toml locally and in Streamlit's Secrets box for the cloud.",
])

H2("Publishing the HTML dashboard (GitHub Pages)")
NUMBERED([
    "Repo Settings → Pages.",
    "Source: 'Deploy from a branch' → Branch: main, folder /docs → Save.",
    "After a minute it serves at your github.io URL.",
])

H2("The real problems we hit (and fixed) — great interview stories")
CALLOUT("1. Hugging Face went paid",
        "Mid-2026 HF required a paid plan to run compute Spaces. Fix: because the LLM was already provider-"
        "swappable, I switched generation to Gemini's free API and deployed on Streamlit's free tier instead.",
        fill="FFF8E1", bar=ORANGE, emoji="🧯")
CALLOUT("2. The model name 404'd",
        "'gemini-2.5-flash' was blocked for new accounts. Fix: I listed the account's available models via the "
        "API and switched the default to the stable alias 'gemini-flash-latest', which won't break when a dated "
        "version retires.",
        fill="FFF8E1", bar=ORANGE, emoji="🧯")
CALLOUT("3. Empty answers / 400 error from the newer model",
        "The newest Flash model rejected a 'no-thinking' setting and could spend its token budget 'thinking', "
        "returning blank text. Fix: made that setting best-effort (retry without it) and raised the output token "
        "limit so the final answer always fits.",
        fill="FFF8E1", bar=ORANGE, emoji="🧯")
CALLOUT("4. Secrets file wouldn't parse",
        "The key pasted without quotes and with a hidden byte-order-mark broke the TOML file. Fix: rewrote it as "
        "valid UTF-8 without BOM and quoted the value. Lesson: config formats are strict.",
        fill="FFF8E1", bar=ORANGE, emoji="🧯")
PAGEBREAK()

# ======================================================================
#  PART 6 — RESULTS & LIMITATIONS
# ======================================================================
H1("📊 Part 6 — Results, Decisions & Honest Limitations", color=PINK)

H2("The headline numbers")
TABLE(["Metric", "Result", "Over"], [
    ["Retrieval recall@5", "1.00", "24 answerable questions"],
    ["Refusal accuracy", "1.00", "26 questions (24 answer + 2 decline)"],
    ["Faithfulness — grounded answer", "0.97", "supported by evidence"],
    ["Faithfulness — hallucinated answer", "0.01", "correctly flagged as unsupported"],
    ["Corpus size", "10 docs · 720 pages · ~1.4M chars · 1,252 chunks", ""],
    ["Answerable score range", "+0.32 to +7.54", "all above the 0.0 threshold"],
    ["Out-of-scope score range", "−8.7 to −9.6", "all safely below threshold"],
], header_fill="C2185B")

H2("Key design decisions & why (defend these)")
TABLE(["Decision", "Why"], [
    ["Hybrid retrieval (not pure semantic)", "Regulation has exact clause numbers/acronyms where keyword search wins."],
    ["Cross-encoder re-ranking", "Far more accurate relevance than vectors; cheap on a 50-item shortlist."],
    ["Reuse rerank score for refusal", "One trustworthy signal for both ranking and the answer/decline decision."],
    ["Threshold in a config file", "Risk teams tune strictness without code changes; tracked in Git."],
    ["Small local LLM", "Extractive cited answers don't need a frontier model; free, offline, private."],
    ["Provider-swappable LLM", "Let us move to Gemini for the free cloud deploy with no rewrite."],
    ["LLM-free metrics in CI", "Deterministic and fast, so the gate is reliable."],
    ["384-token chunks", "Validated by experiment — best balance of citation precision and context."],
], header_fill="C2185B")

H2("Honest limitations (say these before you're asked)")
BULLETS([
    "The corpus is a focused **10-document slice** — the assistant only knows this scope (by design it declines outside it).",
    "Recall is measured at **document level** on a **small, hand-authored** golden set — a baseline, not proof at scale.",
    "The local **3B model** is capable but smaller than hosted frontier models.",
    "Regulations change; citations point to the documents **as downloaded** on a fixed date.",
    "It is a **portfolio demonstration, not legal advice**, and not a production service with real traffic.",
])
CALLOUT("Why stating limits is a strength",
        "In interviews, volunteering limitations signals real understanding and honesty. It's the opposite of an "
        "inflated demo, and it's exactly what senior engineers do.",
        fill="E8F5E9", bar=GREEN, emoji="✅")
PAGEBREAK()
print("part 4+5+6 done")


# ======================================================================
#  PART 7 — INTERVIEW Q&A BANK
# ======================================================================
H1("🎤 Part 7 — Interview Q&A Bank", color=PURPLE)
P("Read the question, try to answer in your head, then read the answer. The **💡 Say this too** lines are "
  "bonus points that make you sound senior. Answers are written the way you'd actually speak.")

H2("A · Project overview & motivation")
QA("Tell me about this project.",
   "I built Grounded, a citation-enforced RAG assistant for banking compliance. It answers questions over 10 "
   "real RBI Master Directions, cites the exact page it used, and refuses to answer when the documents don't "
   "support one — which prevents hallucination. It uses hybrid retrieval, a cross-encoder re-ranker, a citation-"
   "enforcement rule, and a small LangGraph agent, plus a golden evaluation set, faithfulness scoring, and a CI "
   "quality gate. It's deployed as a free public web app.",
   tips="Keep it to ~30 seconds, then pause and let them pick a thread to dig into.")
QA("What problem does it solve?",
   "Bank and NBFC compliance teams answer precise questions against dense, cross-referencing RBI regulation. A "
   "wrong answer carries real regulatory cost, and a normal chatbot will confidently hallucinate. Grounded gives "
   "answers that are traceable to a source page, and it declines rather than guess — which is exactly the "
   "behaviour a compliance team needs.")
QA("Why did you choose RBI regulations as the data?",
   "Three reasons. It's real, public, current data. It ties to my banking background so it's a coherent portfolio "
   "story. And regulatory text genuinely tests the system — it has exact clause numbers and acronyms, which is why "
   "hybrid retrieval is justified, and cross-references and supersession, which stress citation accuracy.")
QA("Who would use this and what's the business value?",
   "Compliance officers, auditors, risk and legal teams, and new-joiner training. The value: faster answers, "
   "audit-ready page citations, much lower hallucination risk, no data leaving the premises (offline option), and "
   "zero per-query cost.",
   tips="Frame value as risk-reduction, not just speed — that's what compliance buyers care about.")
QA("Is this production-ready?",
   "It's a production-grade demonstration of the pattern, not a deployed product. The engineering practices are "
   "production-minded — evaluation, a CI quality gate, versioned prompts, reproducibility — but the corpus is a "
   "focused slice and the golden set is small. I'm deliberately honest about that.")

H2("B · RAG fundamentals")
QA("What is RAG and why use it instead of just an LLM?",
   "RAG is Retrieval-Augmented Generation: you first retrieve relevant text from your own documents, then let the "
   "LLM answer using that text. A plain LLM answers from memory and can hallucinate or be out of date. RAG grounds "
   "answers in real, current sources you control, and lets you cite them.")
QA("Walk me through your RAG pipeline end to end.",
   "Documents are profiled, chunked into 384-token page-tagged pieces, and embedded into 384-dim vectors. At query "
   "time I run hybrid retrieval (BM25 + semantic), fuse with RRF, and re-rank the shortlist with a cross-encoder. A "
   "citation guard checks the top score: if it's too weak, decline; otherwise assemble the cited context and the "
   "LLM writes an answer citing the page. A LangGraph agent orchestrates this and verifies the citation.")
QA("What are the main failure modes of RAG?",
   "Retrieval misses (the right passage isn't fetched), bad chunking (context split or too coarse), the LLM "
   "ignoring the context and using its own memory, and no refusal path so it answers even without evidence. I "
   "address these with hybrid retrieval + re-ranking, a validated chunk size, a strict system prompt, and citation "
   "enforcement with refusal.")
QA("How do you stop the LLM from using outside knowledge?",
   "The system prompt hard-codes 'answer ONLY from the provided context, never use outside knowledge, cite the "
   "page, or return the refusal message'. Temperature is 0.1. And a verify step checks the answer actually cites a "
   "page. It's not a mathematical guarantee, but combined with refusal it strongly constrains the model.")
QA("Why chunk documents? Why not embed whole pages or documents?",
   "Whole documents are too big for the embedder's token limit and too coarse for precise citation — you'd retrieve "
   "a 300-page PDF, not the relevant paragraph. Small chunks give tight, citable evidence. I chunk per page so each "
   "piece keeps its page number.")
QA("How did you choose your chunk size?",
   "By experiment. I tried 256, 384, and 480 tokens. 384 with 64 overlap gave the best balance — 1,252 chunks, "
   "avg ~270 tokens, none over the 512 limit — fragmented enough for precise citation but coherent enough to hold a "
   "complete idea. I chose it with evidence, not a guess.")
QA("What is chunk overlap and why 64 tokens?",
   "Overlap means consecutive chunks share some text so an idea that sits on a boundary isn't cut in half and lost "
   "to retrieval. 64 tokens (about a sixth of the chunk) is enough to preserve boundary context without too much "
   "duplication.")

H2("C · Embeddings & vector search")
QA("What is an embedding?",
   "A list of numbers — here 384 — that represents the meaning of a piece of text. Texts with similar meaning get "
   "vectors pointing in similar directions, so we can find relevant text by comparing vectors instead of matching "
   "exact words.")
QA("Which embedding model did you use and why?",
   "BAAI/bge-small-en-v1.5. It's small, runs on CPU, is fully offline and free, outputs 384-dim vectors, and its "
   "retrieval quality is strong for its size. For extractive retrieval you don't need a huge embedder.")
QA("What is cosine similarity?",
   "A measure of how aligned two vectors are, from −1 to 1. If both are normalized to length 1, it's just their dot "
   "product. 1 means same direction (same meaning), 0 means unrelated. I rank chunks by cosine similarity to the "
   "question.")
QA("Why normalize the embeddings?",
   "bge vectors are designed to be L2-normalized; after that, cosine similarity equals a simple dot product, which "
   "is faster and cleaner. It also keeps all vectors on the same scale.")
QA("Why prepend an instruction to the query but not the passages?",
   "bge-en-v1.5 was trained that way for retrieval: the query gets a short instruction prefix "
   "('Represent this sentence for searching relevant passages') while stored passages are embedded as-is. Matching "
   "the training setup gives better retrieval.")
QA("What is a vector database and why not a normal one?",
   "A vector database indexes embeddings and finds nearest neighbours by meaning-distance, instead of exact matches "
   "like a normal database. I used ChromaDB with cosine distance, storing each chunk's vector, text, and (document, "
   "page) metadata — the metadata is what makes citations possible.")
QA("A classic RAG bug is embedding queries and documents differently — did you avoid it?",
   "Yes. I use the same model for both chunks and queries, and I pass my own embeddings into the store rather than "
   "letting two different components embed with different models. Consistent embedding space is essential.")

H2("D · Retrieval: BM25, hybrid, RRF, re-ranking")
QA("What is BM25?",
   "A classic keyword-ranking formula. It scores a document by how many query words it contains, weighting rare "
   "words more (via IDF) and preventing long documents from dominating. It's excellent at exact matches — clause "
   "numbers, acronyms — which semantic search can miss.")
QA("Why hybrid retrieval instead of just semantic search?",
   "Because regulatory users search exact tokens constantly — a clause number or 'KFS'. Pure semantic search is "
   "great at meaning but weak on exact tokens; BM25 is the opposite. Using both and fusing gets the best of each. "
   "This corpus specifically justifies hybrid.")
QA("What is Reciprocal Rank Fusion and why use it?",
   "RRF merges multiple ranked lists into one using each item's rank position, not its raw score: an item's fused "
   "score is the sum of 1/(k + rank) across lists. I use it because BM25 scores and cosine similarities live on "
   "totally different scales — RRF sidesteps that by using ranks, and it needs almost no tuning.")
QA("What is a cross-encoder and how is it different from the embedder?",
   "The embedder is a bi-encoder: it encodes the query and a passage separately, then compares — fast, good for "
   "scanning everything. A cross-encoder reads the query and passage together in one pass and outputs a single "
   "relevance score — much more accurate but slower, so I only run it on the ~50 shortlisted candidates.")
QA("Why two stages (retrieve then re-rank)?",
   "Speed vs accuracy. The cheap hybrid stage scans all 1,252 chunks and narrows to ~50. The expensive, accurate "
   "cross-encoder then re-ranks just those 50. You get near cross-encoder quality at near first-stage speed.")
QA("How many candidates do you re-rank, and how many do you keep?",
   "I fuse the top ~50 from each retriever, re-rank those, and keep the top few (top-k, e.g. 3–5) as the evidence "
   "passed to the guard and the LLM.")
QA("What does the cross-encoder score actually mean?",
   "It's a relevance logit — higher means more relevant. Positive roughly means 'relevant', negative 'not "
   "relevant', with 0 being the natural 50%% boundary. I reuse this score as the confidence signal for the refusal "
   "decision.")

H2("E · Citation enforcement & refusal")
QA("How does the system decide to refuse?",
   "After re-ranking, I look at the top passage's cross-encoder score. If it's below the threshold (0.0), the "
   "system returns a fixed refusal message instead of answering. Otherwise it builds the cited context and answers.")
QA("Why is the threshold 0.0?",
   "It's the cross-encoder's natural boundary — a logit of 0 corresponds to about 50%% relevance probability. And "
   "it's empirically justified: answerable questions scored from +0.3 up to +7.5, while out-of-scope questions "
   "topped out around −8.7. Zero cleanly separates them with a big margin.")
QA("Why put the threshold and prompts in a config file instead of the code?",
   "So policy can change without touching code and every change is version-controlled in Git. A risk team could "
   "tune strictness. I actually bumped the prompt from v1 to v2 to fix citation formatting with zero code changes — "
   "that's the payoff of externalized config.")
QA("Isn't refusing sometimes just the system failing to find the answer?",
   "Yes, and that's an acceptable trade in compliance. A false 'I don't have that' is far cheaper than a confident "
   "wrong answer. The threshold is tunable if you want to trade refusals for coverage, and my out-of-scope tests "
   "confirm it refuses on genuinely absent topics.")
QA("How do you make citations trustworthy — couldn't the model fake a citation?",
   "Two defenses. The context I hand the model is pre-labelled with the exact bracketed tag for each passage, and "
   "the prompt says to copy that exact tag and never invent one. Then a verify step checks the answer actually "
   "contains a citation in the expected format. If it doesn't, that's flagged.")
print("part 7A-E done")


H2("F · The agent & LangGraph")
QA("What makes this an 'agent' and not just a function?",
   "Two things: it makes a decision (assess the evidence, then branch to answer or decline) rather than following "
   "one fixed path, and it self-checks (a verify step confirms the answer cited a page). The decision branch is the "
   "agentic part. The graph structure also lets me add tools or steps later without rewriting the flow.")
QA("Describe your agent's graph.",
   "It's a LangGraph state machine: START → retrieve → assess → conditional branch to either generate or decline → "
   "(generate) verify → END. Retrieve gets candidates, assess runs the citation guard, generate calls the LLM on "
   "the cited context, verify checks the citation exists, decline returns the refusal.")
QA("Why LangGraph specifically?",
   "It models an agent as an explicit graph with conditional edges, which makes the control flow visible and "
   "auditable — you can point to exactly why it answered or declined. That auditability matters in compliance, and "
   "it's easy to extend with new nodes like query-rewriting.")
QA("What does the verify step do and why?",
   "It confirms the generated answer actually contains a page citation in the expected bracket format. It's a cheap "
   "guard against the model answering without grounding, and it sets a 'verified' flag we surface in the UI.")
QA("Does the LLM run when the system declines?",
   "No — and that's a nice property. For an out-of-scope question the flow stops at 'assess' and returns the "
   "refusal; the LLM is never called. That saves time and cost and guarantees no hallucinated text is even "
   "generated.")

H2("G · The LLM & generation")
QA("Which LLM do you use?",
   "Two, behind one interface. Locally, Qwen2.5-3B-Instruct as a 4-bit quantized GGUF run by llama.cpp on CPU — "
   "offline and free. In the cloud, Google Gemini's free tier. The agent code doesn't change; only the provider "
   "file does.")
QA("Why such a small model?",
   "For grounded, extractive, cited answers the LLM is mostly rephrasing supplied evidence and copying citation "
   "tags — that doesn't need a frontier model. A 3B model is enough, and small means free, offline, and private, "
   "which is exactly what a bank needs.")
QA("What is quantization and why does it matter here?",
   "Quantization stores the model's numbers at lower precision — 4-bit instead of 16-bit — shrinking memory a lot "
   "with minimal quality loss. It's what lets a 3-billion-parameter model run on a normal CPU laptop.")
QA("How is the provider made swappable?",
   "A get_llm() factory returns either a LocalLLM or a GeminiLLM, both exposing the same chat(system, user) method. "
   "It auto-selects Gemini if an API key is present, else local. So switching providers is a one-line/config "
   "change, not a rewrite — which is exactly what saved me when I had to move to the cloud.")
QA("Why temperature 0.1?",
   "Low temperature makes the output focused and near-deterministic. A compliance assistant must be factual and "
   "repeatable, never 'creative'. I don't want different answers to the same question.")
QA("How did you handle the newer Gemini models 'thinking' and returning empty text?",
   "Newer Flash models can spend their output-token budget on internal reasoning and return blank text, and they "
   "rejected the old 'disable thinking' parameter. I made that setting best-effort — try it, and if the model "
   "rejects it, retry without it — and raised the max output tokens so the final answer always fits.")

H2("H · Evaluation & metrics")
QA("How do you know the system actually works?",
   "I measure it. I hand-wrote a golden set of 26 questions — 24 answerable, each mapped to its correct document, "
   "and 2 out-of-scope — and computed retrieval recall@5 and refusal accuracy. Both are 1.00. I also score "
   "faithfulness with an NLI model.")
QA("What is recall@5 and why that metric?",
   "Of the questions that have a correct document, recall@5 is how often that document appears in the top 5 "
   "retrieved results. It directly measures 'did retrieval surface the right source?', which is the foundation of a "
   "correct cited answer.")
QA("What is refusal accuracy?",
   "The fraction of questions where the answer/decline decision was correct — answering in-scope questions and "
   "declining out-of-scope ones. It measures the safety behaviour, which is the whole point of the system.")
QA("Your metrics are 1.00 — isn't that suspicious?",
   "It's a strong baseline, and I'm upfront about why: the golden set is small and hand-authored, and recall is "
   "measured at document level, not exact-page. On a small, clean, in-domain set, perfect recall is plausible. The "
   "honest framing is 'a defensible baseline with a CI gate to catch regressions as the set grows', not 'solved'.",
   tips="Volunteering this caveat unprompted earns big trust.")
QA("Why measure recall at document level and not exact page?",
   "It was the pragmatic first metric to author quickly and reliably. Exact-page recall is on my roadmap — it's "
   "stricter and more informative, but needs more careful ground-truth labelling.")
QA("Why don't you run faithfulness in CI?",
   "Faithfulness needs generated answers, and generation on CPU is slow and non-deterministic — bad for a gate that "
   "must be fast and reliable. So CI runs the deterministic, LLM-free metrics (recall, refusal), and faithfulness "
   "is evaluated locally and in the notebook.")
QA("How would you evaluate this more rigorously with more time?",
   "Expand the golden set to 50–200 questions including adversarial and cross-document ones, add exact-page recall "
   "and precision, add MRR or nDCG for ranking quality, and run faithfulness on a larger sample. I'd also add "
   "confidence intervals given the small set.")

H2("I · Faithfulness / NLI")
QA("What is faithfulness and how do you measure it?",
   "Faithfulness asks: is every claim in the answer actually supported by the retrieved evidence? I measure it with "
   "a Natural Language Inference model that checks, sentence by sentence, whether the evidence entails each claim — "
   "offline and free, no expensive LLM-judge.")
QA("What is NLI?",
   "Natural Language Inference — a model that, given a premise and a hypothesis, decides whether the premise entails "
   "it, is neutral, or contradicts it. I treat evidence sentences as premises and answer sentences as hypotheses, "
   "and use the entailment probability as support.")
QA("You said you fixed a faithfulness bug — what was it?",
   "Initially I fed whole passages as the premise, and the scores came out mushy — a long passage dilutes the "
   "signal. I switched to checking each answer sentence against individual context sentences, which sharpened it. "
   "After that, a grounded answer scored 0.97 and a deliberately fabricated one scored 0.01 — a clean separation.",
   tips="This 'I debugged the metric' story shows genuine depth.")
QA("Why not use GPT-4 as a judge for faithfulness?",
   "Cost, privacy, and reproducibility. An LLM-judge costs money per call, sends data to a third party, and can be "
   "non-deterministic. An offline NLI model is free, private, and repeatable — better for continuous measurement.")

H2("J · CI/CD, deployment & Git")
QA("What is your CI quality gate and why does it matter?",
   "A GitHub Actions workflow runs the golden-set metrics on every push. If recall or refusal drop below 0.85, the "
   "script exits non-zero and the build fails, blocking the change. It's the difference between 'worked once on my "
   "laptop' and quality that's continuously guaranteed.")
QA("Why thresholds of 0.85 when you score 1.00?",
   "Headroom. I don't want the build to fail on tiny, harmless variation, but I do want it to catch a real "
   "regression. 0.85 is comfortably below current performance but well above 'broken'.")
QA("How is the project reproducible?",
   "The CI runs on a clean Linux machine, rebuilds the retriever from the committed PDFs, and runs the tests — "
   "proving it isn't dependent on my machine. I also pin the full environment in requirements-lock.txt.")
QA("How did you deploy the live app?",
   "On Streamlit Community Cloud, free. I point it at dashboard/streamlit_app.py on the main branch and put the "
   "Gemini API key in Streamlit's Secrets. The app loads precomputed embeddings committed in the repo, so cold "
   "starts are fast.")
QA("How do you keep the API key secure?",
   "It never touches the code or Git. Locally it's in a git-ignored .streamlit/secrets.toml; in the cloud it's in "
   "Streamlit's Secrets store. I commit only a .example template so the format is documented without exposing the key.")
QA("Why two requirements files?",
   "requirements.txt is a lean, CPU-only set that installs cleanly on the free cloud host. requirements-lock.txt is "
   "the full pinned environment for exact local reproduction. The cloud doesn't need heavy dev tools or llama.cpp.")

H2("K · Design decisions & trade-offs")
QA("What was your hardest design decision?",
   "Balancing quality against running for free on a laptop with no GPU. I chose small CPU-friendly models, a "
   "two-stage retrieve-then-rerank design for efficiency, and precomputed embeddings — so it's genuinely good while "
   "staying free and offline.")
QA("What would you do differently or improve next?",
   "Grow the golden set with adversarial and cross-document questions, add exact-page recall, and add a query-"
   "rewriting node in the agent for messy questions. I'd also add caching and a small feedback loop for wrong "
   "answers.")
QA("How would this scale to thousands of documents?",
   "Swap the in-memory index for a proper vector database with an approximate-nearest-neighbour index (like "
   "ChromaDB/FAISS at scale), keep BM25 via a search engine, batch the embedding, and possibly add metadata "
   "filtering (by regulation, date) before retrieval. The two-stage design still holds.")
QA("Why not just fine-tune an LLM on the regulations instead of RAG?",
   "Fine-tuning bakes knowledge into weights: it's expensive, hard to update when rules change, gives no citations, "
   "and can still hallucinate. RAG keeps sources external, so answers are current, citable, and the corpus updates "
   "by re-indexing, not re-training.")
QA("How does this generalize beyond RBI?",
   "The pipeline is corpus-agnostic. Swap the documents and re-index and you have the same cite-or-refuse assistant "
   "for insurance policies, legal contracts, medical guidelines, or internal SOPs. Only the data changes.")

H2("L · Troubleshooting stories (show you can debug)")
QA("Tell me about a problem you hit and how you solved it.",
   "My deploy target, Hugging Face Spaces, became paid mid-project. Because I'd built the LLM as a swappable "
   "provider, I switched generation to Gemini's free API and deployed on Streamlit's free tier instead — no "
   "rewrite. Good architecture turned a blocker into a one-file change.")
QA("Give another debugging example.",
   "After moving to Gemini, the default model name 404'd for new accounts and the newest model returned empty "
   "answers because it spent its token budget 'thinking'. I listed the account's available models via the API, "
   "switched to the stable 'gemini-flash-latest' alias, made the thinking setting best-effort, and raised the "
   "output-token limit. I validated it end-to-end before deploying.")
QA("How do you debug a wrong or missing retrieval?",
   "I inspect the ranked candidates and their scores at each stage — BM25, semantic, fused, and cross-encoder. That "
   "shows whether the right passage was never retrieved (a recall/chunking issue) or retrieved but ranked low (a "
   "re-ranking issue), which points to the fix.")

H2("M · Behavioural / portfolio framing")
QA("What are you most proud of in this project?",
   "That it's honest. It refuses when it should, it's measured with real metrics and stated limitations, and every "
   "answer is traceable to a page. It's engineered like something a team could trust, not a flashy demo.")
QA("What did you learn?",
   "How much of production RAG is about trust and evaluation, not just wiring an LLM: refusal, citations, a golden "
   "set, faithfulness, and a CI gate. I also learned the value of swappable design — it saved my deployment.")
QA("How long did it take and how did you work?",
   "I built it in phases — a working pipeline first, then production quality, then the agent, evaluation, and "
   "deployment — committing each phase to Git with clear messages and validating each step before moving on.")

H2("N · Curveballs & advanced")
QA("What if two regulations contradict each other?",
   "Currently it retrieves and cites the strongest-matching passages, so a user sees both sources with pages and "
   "can judge. A stronger version would detect conflicting evidence and surface it explicitly — that's a good "
   "roadmap item, and the agent graph makes adding such a node straightforward.")
QA("How would you add multi-hop questions (answer spans several documents)?",
   "Add a query-decomposition or iterative-retrieval node in the agent: retrieve, see what's missing, issue a "
   "follow-up query, and combine evidence before answering. LangGraph is designed for exactly this kind of "
   "multi-step flow.")
QA("What are the security or privacy concerns?",
   "The main ones are data leaving the premises and secret leakage. The offline local-LLM mode keeps all data on "
   "the machine; the API key is never committed. For a real deployment I'd add access controls and audit logging.")
QA("Could a user prompt-inject the system into ignoring the rules?",
   "It's a real risk with any LLM. Mitigations here: the system prompt is strict, the model only sees retrieved "
   "regulation text plus the question, temperature is low, and the verify step checks for a citation. For "
   "production I'd add input sanitisation and output validation, and treat retrieved text as untrusted too.")
QA("Why not use an existing framework like LlamaIndex end-to-end?",
   "I wanted to understand and control each stage — chunking, hybrid fusion, the refusal rule, the agent — for a "
   "portfolio piece where I can defend every choice. I did use focused libraries (sentence-transformers, rank-bm25, "
   "LangGraph) rather than reinventing them, but kept the pipeline logic explicit.")
QA("What's the difference between your system declining and just having low confidence?",
   "They're the same signal here: the cross-encoder relevance score IS the confidence, and the threshold turns low "
   "confidence into an explicit refusal. That's deliberate — one interpretable number drives both ranking and the "
   "decline decision.")
QA("If retrieval recall is perfect, why do you even need the re-ranker?",
   "Recall@5 being perfect means the right document is in the top 5, but the re-ranker decides which passage is "
   "number one and provides the confidence score that powers refusal. Good ranking and a calibrated score matter "
   "even when recall is high — and recall won't stay perfect as the corpus grows.")
QA("What happens if the embedding model and corpus drift over time?",
   "If I change the embedding model I must re-embed the whole corpus so queries and chunks share one space. If "
   "regulations change, I re-index the new PDFs. The CI gate and golden set catch quality regressions from either.")
PAGEBREAK()

# ======================================================================
#  PART 8 — CHEAT SHEET
# ======================================================================
H1("⚡ Part 8 — Night-Before Cheat Sheet", color=NAVY)
P("If you remember nothing else, remember this page.")

CALLOUT("One-liner",
        "A citation-enforced RAG assistant over real RBI regulations: hybrid retrieval + cross-encoder re-ranking "
        "+ a refusal rule + a LangGraph agent; cites the exact page or honestly declines. Free & deployable.",
        fill="E8F5E9", bar=GREEN, emoji="⭐")

H2("The numbers")
BULLETS([
    "**10** RBI documents · **720** pages · **~1.4M** characters · **1,252** chunks",
    "Chunks: **384 tokens**, **64** overlap, page-tagged",
    "Embeddings: **bge-small-en-v1.5**, **384** dimensions, on CPU",
    "Re-ranker: **ms-marco-MiniLM** cross-encoder; refusal threshold **0.0**",
    "Golden set: **26** questions (24 answer + 2 decline)",
    "**Recall@5 = 1.00** · **Refusal accuracy = 1.00** · **Faithfulness 0.97 vs 0.01**",
    "CI thresholds: recall & refusal ≥ **0.85**",
    "LLM: **Qwen2.5-3B** (local, offline) or **Gemini** (cloud, free)",
])

H2("The flow (say it in order)")
P("**profile → chunk → embed → store → hybrid retrieve (BM25 + semantic, RRF) → cross-encoder re-rank → "
  "citation guard (cite or refuse) → LangGraph agent → LLM writes cited answer → verify citation**")

H2("The five ideas that impress")
BULLETS([
    "**Refusal is the feature** — declining beats a confident wrong answer in compliance.",
    "**Hybrid retrieval is justified by the data** — regulation has exact clause numbers.",
    "**The re-rank score doubles as the confidence signal** for refusal.",
    "**Provider-swappable LLM** — the design that let me deploy free when hosting changed.",
    "**Measured, not claimed** — golden set, faithfulness, and a CI gate; limitations stated honestly.",
])

CALLOUT("If you get stuck in an interview",
        "Say: \"Let me reason through it.\" Then narrate the flow above. Almost every question maps to one step of "
        "that pipeline — name the step, say what it does and why, and you're answering well.",
        fill="E3F2FD", bar=BLUE, emoji="🧠")

P("")
P("— End of the Grounded Project Bible —", align="center", color=GREY, italic=True)

# ======================================================================
#  SAVE
# ======================================================================
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print("SAVED:", OUT)
